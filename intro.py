from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)

document = Document()

name = 'Alibaba Abdulbasit Olamide'
phone_number = '08163964660'
email = 'abdulbasitalibaba@gmail.com'

speak('Hello' + name + "How are you doing?")

# profile picture

document.add_picture(
    'profile.png',
    width=Inches(1.0))

# details
document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About me')
document.add_paragraph("I'm a data analyst from nigeria")

document.add_heading('Work experience')
para = document.add_paragraph()

company = input("where did you worked")
period = input('when did you worked there')

para.add_run(company + ' ').bold = True
para.add_run(period + '\n')

experience = input('describe your experience')
para.add_run(experience)

# more experience
while True:
    more_experience = input("do you have more experience ")
    if more_experience.lower() == 'yes':

        para = document.add_paragraph()

        company = input("where did you worked? ")
        period = input('when did you worked there? ')

        para.add_run(company + ' ').bold = True
        para.add_run(period + '\n')

        experience = input('describe your experience ')
        para.add_run(experience)
    else:
        break

# skills section

document.add_heading("skills")

skill = input('What skill do you have ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    more_skill = input('do you have more skills? ')
    if more_skill.lower() == 'beeni':
        skill = input('What skill do you have more? ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

    
document.save('cv.docx')